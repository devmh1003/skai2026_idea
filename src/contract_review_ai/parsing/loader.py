"""계약서 파일 → 평문 텍스트 → Document.

지원 형식
    .txt .md   — 그대로
    .hwpx      — OWPML(zip+XML). 표준 라이브러리만으로 파싱한다.
    .hwp       — 구형 한글 바이너리. olefile이 있으면 best-effort로 추출한다.
    .docx      — python-docx (표 포함)
    .pdf       — pypdf
"""

from __future__ import annotations

import re
import zipfile
from pathlib import Path
from xml.etree import ElementTree

from ..models import Document
from ..parties import detect_parties
from .segmenter import segment_clauses

TEXT_SUFFIXES = {".txt", ".md", ".text"}
SUPPORTED_SUFFIXES = TEXT_SUFFIXES | {".hwpx", ".hwp", ".docx", ".pdf"}


def read_text(path: str | Path) -> str:
    p = Path(path)
    if not p.is_file():
        raise FileNotFoundError(f"계약서 파일을 찾을 수 없습니다: {p}")

    suffix = p.suffix.lower()
    if suffix in TEXT_SUFFIXES or suffix == "":
        return p.read_text(encoding="utf-8-sig")
    if suffix == ".hwpx":
        return read_hwpx(p)
    if suffix == ".hwp":
        return read_hwp(p)
    if suffix == ".docx":
        return _read_docx(p)
    if suffix == ".pdf":
        return _read_pdf(p)
    raise ValueError(
        f"지원하지 않는 형식입니다: {suffix} "
        f"(지원: {' '.join(sorted(SUPPORTED_SUFFIXES))})"
    )


# ---------------------------------------------------------------- 한글(HWPX)

_SECTION_RE = re.compile(r"Contents/section(\d+)\.xml$", re.IGNORECASE)


def read_hwpx(path: Path) -> str:
    """HWPX(OWPML)에서 문단 텍스트를 뽑는다.

    구조: zip 안의 Contents/section0.xml, section1.xml … 각 섹션이 문단(hp:p)을
    담고, 문단 안의 hp:t 요소에 실제 글자가 들어 있다. 표(hp:tbl) 안의 문단도
    같은 hp:p 구조라 별도 처리 없이 순서대로 읽힌다.
    """
    try:
        archive = zipfile.ZipFile(path)
    except zipfile.BadZipFile as exc:
        raise ValueError(
            f"HWPX 파일을 열 수 없습니다(구형 .hwp를 .hwpx로 저장했는지 확인): {path}"
        ) from exc

    with archive:
        sections = sorted(
            (n for n in archive.namelist() if _SECTION_RE.search(n)),
            key=lambda n: int(_SECTION_RE.search(n).group(1)),  # type: ignore[union-attr]
        )
        if not sections:
            raise ValueError(f"HWPX 본문(Contents/sectionN.xml)을 찾지 못했습니다: {path}")

        paragraphs: list[str] = []
        for name in sections:
            root = ElementTree.fromstring(archive.read(name))
            paragraphs.extend(_hwpx_paragraphs(root))

    return "\n".join(paragraphs)


def _hwpx_paragraphs(root: ElementTree.Element) -> list[str]:
    paragraphs: list[str] = []
    for element in root.iter():
        if _local(element.tag) != "p":
            continue
        runs = [
            child.text
            for child in element.iter()
            if _local(child.tag) == "t" and child.text
        ]
        text = "".join(runs).strip()
        # 표 안의 문단이 바깥 문단에 중복으로 잡히지 않도록 빈 문단은 버린다.
        if text:
            paragraphs.append(text)
    return paragraphs


def _local(tag: str) -> str:
    return tag.rsplit("}", 1)[-1]


# ---------------------------------------------------------------- 한글(구형 HWP)


def read_hwp(path: Path) -> str:
    """구형 .hwp(HWP 5.0, CFB 컨테이너)에서 문단 텍스트를 best-effort로 추출한다.

    BodyText/SectionN 스트림을 (필요시 zlib 해제 후) 레코드 단위로 읽어
    HWPTAG_PARA_TEXT(67) 레코드의 UTF-16LE 본문만 취한다. 표·글상자 등 일부
    개체의 텍스트는 누락될 수 있으므로, 정확한 검토가 필요하면 hwpx나 docx로
    저장해 다시 넣는 편을 권한다.
    """
    try:
        import olefile  # type: ignore[import-not-found]
    except ImportError as exc:  # pragma: no cover - 선택 의존성
        raise RuntimeError(
            "구형 .hwp를 읽으려면 `pip install olefile`이 필요합니다. "
            "또는 한글에서 .hwpx/.docx/.pdf로 저장해 사용하십시오."
        ) from exc

    import struct
    import zlib

    ole = olefile.OleFileIO(str(path))
    try:
        compressed = _hwp_is_compressed(ole)
        streams = sorted(
            (e for e in ole.listdir() if len(e) == 2 and e[0] == "BodyText"),
            key=lambda e: int("".join(ch for ch in e[1] if ch.isdigit()) or 0),
        )
        if not streams:
            raise ValueError(f"HWP 본문(BodyText)을 찾지 못했습니다: {path}")

        paragraphs: list[str] = []
        for entry in streams:
            data = ole.openstream(entry).read()
            if compressed:
                data = zlib.decompress(data, -15)
            paragraphs.extend(_hwp_paragraphs(data, struct))
    finally:
        ole.close()

    return "\n".join(paragraphs)


_HWPTAG_PARA_TEXT = 67
_HWP_CONTROL_CHARS = set(range(0, 32)) - {9, 10, 13}


def _hwp_is_compressed(ole) -> bool:
    if not ole.exists("FileHeader"):
        return True
    header = ole.openstream("FileHeader").read()
    return bool(header[36] & 0x01) if len(header) > 36 else True


def _hwp_paragraphs(data: bytes, struct) -> list[str]:
    paragraphs: list[str] = []
    cursor = 0
    size = len(data)

    while cursor + 4 <= size:
        (header,) = struct.unpack("<I", data[cursor : cursor + 4])
        tag_id = header & 0x3FF
        length = (header >> 20) & 0xFFF
        cursor += 4
        if length == 0xFFF:  # 확장 길이
            if cursor + 4 > size:
                break
            (length,) = struct.unpack("<I", data[cursor : cursor + 4])
            cursor += 4
        payload = data[cursor : cursor + length]
        cursor += length

        if tag_id != _HWPTAG_PARA_TEXT:
            continue
        text = payload.decode("utf-16-le", errors="ignore")
        cleaned = "".join(
            ch for ch in text if ord(ch) not in _HWP_CONTROL_CHARS
        ).strip()
        if cleaned:
            paragraphs.append(cleaned)

    return paragraphs


# ---------------------------------------------------------------- docx / pdf


def _read_docx(path: Path) -> str:
    try:
        import docx  # type: ignore[import-not-found]
    except ImportError as exc:  # pragma: no cover - 선택 의존성
        raise RuntimeError("docx를 읽으려면 `pip install python-docx`가 필요합니다.") from exc

    document = docx.Document(str(path))
    parts = [para.text for para in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text.strip() for cell in row.cells))
    return "\n".join(parts)


def _read_pdf(path: Path) -> str:
    try:
        from pypdf import PdfReader  # type: ignore[import-not-found]
    except ImportError as exc:  # pragma: no cover - 선택 의존성
        raise RuntimeError("pdf를 읽으려면 `pip install pypdf`가 필요합니다.") from exc

    reader = PdfReader(str(path))
    pages = [page.extract_text() or "" for page in reader.pages]
    text = "\n".join(pages)
    if not text.strip():
        raise ValueError(
            f"PDF에서 텍스트를 추출하지 못했습니다(스캔본일 수 있음): {path}. "
            "OCR을 거치거나 원본 hwpx/docx를 사용하십시오."
        )
    return text


# ---------------------------------------------------------------- 진입점


def load_document(path: str | Path, name: str | None = None, version: str = "") -> Document:
    p = Path(path)
    text = read_text(p)
    clauses = segment_clauses(text)
    if not clauses:
        raise ValueError(f"조문을 하나도 추출하지 못했습니다: {p}")
    return Document(
        name=name or p.stem,
        path=str(p),
        clauses=clauses,
        parties=detect_parties(text),
        version=version,
    )
